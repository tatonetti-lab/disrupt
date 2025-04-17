"""
matcher.py

VERION: 0.1

INPUTS
======
- cancer patients in SEC
- cancer trials in SEC

OUTPUTS
=======
- matched patient and trial list

STEPS
=====

1. Query the datastores
2. Print out the results



"""

import sqlite3
import csv
import os
from datetime import datetime
import argparse
import docx

parser= argparse.ArgumentParser()

parser.add_argument("--disease",type=str, required=True,help="Please specify a diseae site (breast/liver/prostate/bladder)")

args = parser.parse_args()

disease = args.disease

sqlite = sqlite3.connect('disrupt.db')

sqlite_cursor = sqlite.cursor()

today = datetime.today().strftime("%Y-%m-%d")


#NOTE: currently the drug join is left because we found no matches there

#TO-DO: add filter on disease subquery's date of screening in the overall where clause (otherwise we return everything every time) DONE

#TO-DO: ensure ALL receptor statuses match.

if(disease == 'Breast' or disease == 'breast'):

    sql = """select distinct * from (select disease.nci_number,date_screened, new_or_progressed,disease.mrn, cancer_type, stage_match,receptor_match,treatment_match
     from
     (
      select nci_number, mrn,a.pk_id as pt_pk, c.pk_id as trial_pk, a.cancer_type,date_screened, new_or_progressed
     from
     patient a left join
     trial_cancer_type b on a.cancer_type = b.cancer_type left join
     trial c on b.fk_id = c.pk_id
     ) disease
     inner join
     (
     select nci_number, mrn,a.pk_id as pt_pk, d.pk_id as trial_pk, group_concat(distinct c.stage) as stage_match, count(distinct c.stage) as num_matches
     from
     patient a left join
     patient_staging b on a.pk_id = b.fk_id left join
     trial_stage c on  b.stage = c.stage left join
     trial d on c.fk_id = d.pk_id left join
     staging_rank e on b.stage = e.stage_name
     where ranking = (select max(ranking) from staging_rank q inner join patient_staging qq on q.stage_name = qq.stage where qq.fk_id = b.fk_id)
     group by  nci_number, mrn,a.pk_id, d.pk_id
     ) stage on disease.pt_pk = stage.pt_pk and disease.trial_pk = stage.trial_pk
     inner join
     (
    select * from
     (
     select nci_number, mrn, trial_pk, pt_pk,group_concat(distinct case when patient_receptor = trial_receptor and cnt_types_for_this_trial = 1 then patient_receptor else null end) as receptor_match,
     sum(case when receptor_type = 'PR' and (cnt_types_for_this_trial in (0,2) or (trial_receptor = patient_receptor)) then 1 else 0 end) as pr_match,
     sum(case when receptor_type = 'ER' and (cnt_types_for_this_trial in (0,2) or (trial_receptor = patient_receptor)) then 1 else 0 end) as er_match,
     sum(case when receptor_type = 'HER2' and (cnt_types_for_this_trial in (0,2) or (trial_receptor = patient_receptor)) then 1 else 0 end) as her2_match,
     sum(case when receptor_type = 'PR' then 1 else 0 end) as pr_req,
     sum(case when receptor_type = 'ER' then 1 else 0 end) as er_req,
     sum(case when receptor_type = 'HER2' then 1 else 0 end) as her2_req
     from
     (
        select nci_number, mrn,a.pk_id as pt_pk, d.pk_id as trial_pk, c.receptor_type, c.receptor_value,
        (select count(q.receptor_value) from trial_receptor q where c.receptor_type = q.receptor_type and d.pk_id = q.fk_id) as cnt_types_for_this_trial, b.receptor_value as patient_receptor,
        c.receptor_value as trial_receptor
     from
     patient a left join
     patient_receptor b on a.pk_id = b.fk_id left join
     trial_receptor c on  b.receptor_type = c.receptor_type and b.receptor_value = c.receptor_value left join
     trial d on c.fk_id = d.pk_id
    ) group by nci_number, mrn, trial_pk,pt_pk
    ) q where (pr_req = 0 or pr_match>0) and (er_req = 0 or er_match>0) and (her2_req=0 or her2_match>0)
     ) receptor on disease.pt_pk = receptor.pt_pk and disease.trial_pk = receptor.trial_pk
     left join
     (
       select nci_number, mrn,a.pk_id as pt_pk, d.pk_id as trial_pk, group_concat(distinct c.treatment_name) as treatment_match, count(distinct c.treatment_name) as num_matches
     from
     patient a left join
     patient_treatment b on a.pk_id = b.fk_id left join
     trial_treatment c on  b.treatment_name = c.treatment_name and b.treatment_type = c.treatment_type left join
     trial d on c.fk_id = d.pk_id
     group by  nci_number, mrn,a.pk_id, d.pk_id
     ) tx on disease.pt_pk = tx.pt_pk and disease.trial_pk = tx.trial_pk
     where disease.date_screened >= '%s'
    union all
    select disease.nci_number,date_screened, new_or_progressed, disease.mrn, cancer_type, 'ALL STAGES ALLOWED' as stage_match,receptor_match, treatment_match as treatment_match
     from
     (
      select nci_number, mrn,a.pk_id as pt_pk, c.pk_id as trial_pk, a.cancer_type,date_screened, new_or_progressed
     from
     patient a left join
     trial_cancer_type b on a.cancer_type = b.cancer_type left join
     trial c on b.fk_id = c.pk_id
     ) disease
     inner join
     (
    select * from
     (
     select nci_number, mrn, trial_pk, pt_pk,group_concat(distinct case when patient_receptor = trial_receptor and cnt_types_for_this_trial=1 then patient_receptor else null end) as receptor_match,
     sum(case when receptor_type = 'PR' and (cnt_types_for_this_trial in (0,2) or (trial_receptor = patient_receptor)) then 1 else 0 end) as pr_match,
     sum(case when receptor_type = 'ER' and (cnt_types_for_this_trial in (0,2) or (trial_receptor = patient_receptor)) then 1 else 0 end) as er_match,
     sum(case when receptor_type = 'HER2' and (cnt_types_for_this_trial in (0,2) or (trial_receptor = patient_receptor)) then 1 else 0 end) as her2_match,
     sum(case when receptor_type = 'PR' then 1 else 0 end) as pr_req,
     sum(case when receptor_type = 'ER' then 1 else 0 end) as er_req,
     sum(case when receptor_type = 'HER2' then 1 else 0 end) as her2_req
     from
     (
        select nci_number, mrn,a.pk_id as pt_pk, d.pk_id as trial_pk, c.receptor_type, c.receptor_value,
        (select count(q.receptor_value) from trial_receptor q where c.receptor_type = q.receptor_type and d.pk_id = q.fk_id) as cnt_types_for_this_trial, b.receptor_value as patient_receptor,
        c.receptor_value as trial_receptor
     from
     patient a left join
     patient_receptor b on a.pk_id = b.fk_id left join
     trial_receptor c on  b.receptor_type = c.receptor_type and b.receptor_value = c.receptor_value left join
     trial d on c.fk_id = d.pk_id
    ) group by nci_number, mrn, trial_pk,pt_pk
    ) q where (pr_req = 0 or pr_match>0) and (er_req = 0 or er_match>0) and (her2_req=0 or her2_match>0)
     ) receptor on disease.pt_pk = receptor.pt_pk and disease.trial_pk = receptor.trial_pk
     left join
     (
       select nci_number, mrn,a.pk_id as pt_pk, d.pk_id as trial_pk, group_concat(distinct c.treatment_name) as treatment_match, count(distinct c.treatment_name) as num_matches
     from
     patient a left join
     patient_treatment b on a.pk_id = b.fk_id left join
     trial_treatment c on  b.treatment_name = c.treatment_name and b.treatment_type = c.treatment_type left join
     trial d on c.fk_id = d.pk_id
     group by  nci_number, mrn,a.pk_id, d.pk_id
     ) tx on disease.pt_pk = tx.pt_pk and disease.trial_pk = tx.trial_pk
     where not exists (select 1 from trial_stage q where disease.trial_pk = q.fk_id)
     and disease.date_screened >= '%s'
     union all
     select disease.nci_number,date_screened, new_or_progressed,disease.mrn, cancer_type, stage_match,'ALL RECEPTORS ALLOWED' as receptor_match,treatment_match
     from
     (
      select nci_number, mrn,a.pk_id as pt_pk, c.pk_id as trial_pk, a.cancer_type,date_screened,new_or_progressed
     from
     patient a left join
     trial_cancer_type b on a.cancer_type = b.cancer_type left join
     trial c on b.fk_id = c.pk_id
     ) disease
     inner join
     (
     select nci_number, mrn,a.pk_id as pt_pk, d.pk_id as trial_pk, group_concat(distinct c.stage) as stage_match, count(distinct c.stage) as num_matches
     from
     patient a left join
     patient_staging b on a.pk_id = b.fk_id left join
     trial_stage c on  b.stage = c.stage left join
     trial d on c.fk_id = d.pk_id left join
     staging_rank e on b.stage = e.stage_name
     where ranking = (select max(ranking) from staging_rank q inner join patient_staging qq on q.stage_name = qq.stage where qq.fk_id = b.fk_id)
     group by  nci_number, mrn,a.pk_id, d.pk_id
     ) stage on disease.pt_pk = stage.pt_pk and disease.trial_pk = stage.trial_pk
     left join
     (
       select nci_number, mrn,a.pk_id as pt_pk, d.pk_id as trial_pk, group_concat(distinct c.treatment_name) as treatment_match, count(distinct c.treatment_name) as num_matches
     from
     patient a left join
     patient_treatment b on a.pk_id = b.fk_id left join
     trial_treatment c on  b.treatment_name = c.treatment_name and b.treatment_type = c.treatment_type left join
     trial d on c.fk_id = d.pk_id
     group by  nci_number, mrn,a.pk_id, d.pk_id
     ) tx on disease.pt_pk = tx.pt_pk and disease.trial_pk = tx.trial_pk
      where not exists (select 1 from trial_receptor q where disease.trial_pk = q.fk_id) and  disease.date_screened > '%s'
    union all
    select disease.nci_number,date_screened, new_or_progressed,disease.mrn, cancer_type, 'ALL STAGES ALLOWED' as staging_match,'ALL RECEPTORS ALLOWED' as receptor_match,treatment_match
     from
     (
      select nci_number, mrn,a.pk_id as pt_pk, c.pk_id as trial_pk, a.cancer_type,date_screened,new_or_progressed
     from
     patient a left join
     trial_cancer_type b on a.cancer_type = b.cancer_type left join
     trial c on b.fk_id = c.pk_id
     ) disease
     left join
     (
       select nci_number, mrn,a.pk_id as pt_pk, d.pk_id as trial_pk, group_concat(distinct c.treatment_name) as treatment_match, count(distinct c.treatment_name) as num_matches
     from
     patient a left join
     patient_treatment b on a.pk_id = b.fk_id left join
     trial_treatment c on  b.treatment_name = c.treatment_name and b.treatment_type = c.treatment_type left join
     trial d on c.fk_id = d.pk_id
     group by  nci_number, mrn,a.pk_id, d.pk_id
     ) tx on disease.pt_pk = tx.pt_pk and disease.trial_pk = tx.trial_pk 
     where not exists (select 1 from trial_receptor q where disease.trial_pk = q.fk_id)
    and not exists (select 1 from trial_stage q where disease.trial_pk = q.fk_id) and disease.date_screened >= '%s' ) q""" % (today,today,today,today)

elif (disease == 'Lung' or disease == 'lung'):
    sql = """


select 
pt_number, type, therapy_type, mrn,pat_name, nci_number, nct_number, early_stage_resectable, early_stage_unresectable, advanced_first_line,
advanced_second_line, title,spanish_title, english_description, spanish_description,disease_setting,
trial_keyword1, trial_keyword_2, group_concat(patient_gene,',') as pt_genes, pt_stage, trial_stages, date_screened
 from
(

select 
dense_rank () over(order by mrn) as pt_number,
q.* from
(

select * from
(
select 
'1' as type,therapy_type,date_screened,
mrn,a.nci_number, a.nct_number ,
early_stage_resectable, early_stage_unresectable, advanced_first_line, advanced_second_line,
title , spanish_title,
english_description, spanish_description, keyword1 as trial_keyword1, keyword2 as trial_keyword_2, hugo_gene as patient_gene,
(select stage from patient_staging q inner join staging_rank qq on stage=stage_name where q.fk_id = e.pk_id order by ranking desc limit 1) as pt_stage,
(select group_concat(stage) from trial_stage q where a.pk_id = q.fk_id) as trial_stages,disease_setting,pat_name
from
trial a inner join
trial_manual_classification b on a.nci_number = b.nci_number inner join
trial_manual_matchcriteria c on a.nct_number = c.nci_number inner join
patient_genes d on d.hugo_gene like '%' || keyword1 || '%' and (keyword2 is null or d.hugo_gene like '%' || keyword2 || '%') inner join
patient e on d.fk_id = e.pk_id left join
trial_patient_descriptions f on a.nci_number = f.nci_number left join
(select receptor_value as disease_setting, fk_id from patient_receptor where receptor_type = 'DISEASE_SETTING') g on e.pk_id = 
g.fk_id
where therapy_type like '%Targeted%' and
(
        disease_setting is null or
        (disease_setting in ('localized','locally advanced resectable') and early_stage_resectable = 'Y') or
        (disease_setting = 'locally advanced unresectable' and early_stage_unresectable = 'Y') or
        (disease_setting = 'metastatic, first line' and advanced_first_line = 'Y' ) or 
        (disease_setting = 'metastatic, later line' and advanced_second_line = 'Y')
)
) q

union

select * from
(
select distinct case when therapy_type like '%antibody%' then 2 else 3 end as type, therapy_type,date_screened,
mrn,
a.nci_number, a.nct_number,
early_stage_resectable, early_stage_unresectable, advanced_first_line, advanced_second_line,
title, spanish_title, english_description, spanish_description, 'N/A' as trial_keyword1, 'N/A' as trial_keyword_2, 
(select group_concat(distinct hugo_gene) from patient_genes q where d.pk_id = q.fk_id) as patient_gene,
(select stage from patient_staging q inner join staging_rank qq on stage=stage_name where q.fk_id = d.pk_id order by ranking desc limit 1) as pt_stage,
(select group_concat(stage) from trial_stage q where a.pk_id = q.fk_id) as trial_stages, disease_setting,pat_name
from
trial a inner join
trial_manual_classification b on a.nci_number = b.nci_number  join 
patient d left join
trial_patient_descriptions f on a.nci_number = f.nci_number left join
(select receptor_value as disease_setting, fk_id from patient_receptor where receptor_type = 'DISEASE_SETTING') g on d.pk_id = 
g.fk_id
where
therapy_type not like '%Targeted%'
) where (patient_gene like '%EGFR%' or patient_gene like '%ALK%' or patient_gene like '%ROS1%' or  patient_gene like '%RET%' or  patient_gene like '%HER2%' 
or patient_gene like '%ERBB2%' or patient_gene like '%NTRK%' ) 
and
(
        disease_setting is null or
        (disease_setting in ('localized','locally advanced resectable') and early_stage_resectable = 'Y') or
        (disease_setting = 'locally advanced unresectable' and early_stage_unresectable = 'Y') or
        (disease_setting = 'metastatic, first line' and advanced_first_line = 'Y' ) or 
        (disease_setting = 'metastatic, later line' and advanced_second_line = 'Y')
)

union

select * from
(
select distinct case when therapy_type like '%immuno%' then 2  when therapy_type like '%antibody%' then 3 else 4 end as type, therapy_type,date_screened,
mrn,
a.nci_number, a.nct_number,
early_stage_resectable, early_stage_unresectable, advanced_first_line, advanced_second_line,
title,  spanish_title, english_description, spanish_description,'N/A' as trial_keyword1, 'N/A' as trial_keyword_2, 
(select group_concat(distinct hugo_gene) from patient_genes q where d.pk_id = q.fk_id) as patient_gene,
(select stage from patient_staging q inner join staging_rank qq on stage=stage_name where q.fk_id = d.pk_id order by ranking desc limit 1) as pt_stage,
(select group_concat(stage) from trial_stage q where a.pk_id = q.fk_id) as trial_stages,disease_setting,pat_name
from
trial a inner join
trial_manual_classification b on a.nci_number = b.nci_number  join 
patient d left join
trial_patient_descriptions f on a.nci_number = f.nci_number left join
(select receptor_value as disease_setting, fk_id from patient_receptor where receptor_type = 'DISEASE_SETTING') g on d.pk_id = 
g.fk_id
where
therapy_type not like '%Targeted%'
) where not (patient_gene like '%EGFR%' or patient_gene like '%ALK%' or patient_gene like '%ROS1%' or  patient_gene like '%RET%' or  patient_gene like '%HER2%' 
or patient_gene like '%ERBB2%' or patient_gene like '%NTRK%' )  
and
(
        disease_setting is null or
        (disease_setting in ('localized','locally advanced resectable') and early_stage_resectable = 'Y') or
        (disease_setting = 'locally advanced unresectable' and early_stage_unresectable = 'Y') or
        (disease_setting = 'metastatic, first line' and advanced_first_line = 'Y' ) or 
        (disease_setting = 'metastatic, later line' and advanced_second_line = 'Y')
)


union

select * from
(
select distinct case when therapy_type like '%immuno%' then 1  when therapy_type like '%antibody%' then 2 else 3 end as type, therapy_type,date_screened,
mrn,
a.nci_number, a.nct_number,
early_stage_resectable, early_stage_unresectable, advanced_first_line, advanced_second_line,
title, spanish_title, english_description, spanish_description, 'N/A' as trial_keyword1, 'N/A' as trial_keyword_2, 
(select group_concat(distinct hugo_gene) from patient_genes q where d.pk_id = q.fk_id) as patient_gene,
(select stage from patient_staging q inner join staging_rank qq on stage=stage_name where q.fk_id = d.pk_id order by ranking desc limit 1) as pt_stage,
(select group_concat(stage) from trial_stage q where a.pk_id = q.fk_id) as trial_stages,disease_setting,pat_name
from
trial a inner join
trial_manual_classification b on a.nci_number = b.nci_number  join 
patient d left join
trial_patient_descriptions f on a.nci_number = f.nci_number left join
(select receptor_value as disease_setting, fk_id from patient_receptor where receptor_type = 'DISEASE_SETTING') g on d.pk_id = 
g.fk_id
where
therapy_type not like '%Targeted%'
) where patient_gene is null  and
(
        disease_setting is null or
        (disease_setting in ('localized','locally advanced resectable') and early_stage_resectable = 'Y') or
        (disease_setting = 'locally advanced unresectable' and early_stage_unresectable = 'Y') or
        (disease_setting = 'metastatic, first line' and advanced_first_line = 'Y' ) or 
        (disease_setting = 'metastatic, later line' and advanced_second_line = 'Y')
)


) q 
)q
where date_screened>='2024-12-20'
group by 
pt_number, type, therapy_type, mrn, nci_number, nct_number, early_stage_resectable, early_stage_unresectable, advanced_first_line,
advanced_second_line, title,
trial_keyword1, trial_keyword_2, pt_stage, trial_stages,date_screened,spanish_title, english_description, spanish_description,disease_setting,pat_name
ORDER BY 1, CAST(TYPE as integer), nci_number

"""

cursor = sqlite.cursor()

print(sql)

cursor.execute(sql)
results = cursor.fetchall()
today = datetime.today()
print(today.strftime("%Y/%m/%d %H:%M:%S"))

outfile = open("matches/matches_" + today.strftime("%Y-%m-%d") + ".txt", 'w')
writer =csv.writer(outfile)


#pt_number, type, therapy_type, mrn, nci_number, nct_number, early_stage_resectable, early_stage_unresectable, advanced_first_line,advanced_second_line, title, trial_keyword1, trial_keyword_2, pt_genes, pt_stage, trial_stages, date_screened


writer.writerow(['pt_number', 'type', 'therapy_type', 'mrn','patient_name', 'nci_number', 'nct_number', 'early_stage_resectable', 'early_stage_unresectable', 'advanced_first_line','advanced_second_line', 'title','spanish title','english description','spanish description', 'trial_keyword1', 'trial_keyword_2', 'pt_genes', 'pt_stage', 'trial_stages', 'date_screened','disease setting'])

for pt_number, type, therapy_type, mrn,pat_name, nci_number, nct_number, early_stage_resectable, early_stage_unresectable, advanced_first_line,advanced_second_line, title, spanish_title, english_description, spanish_description,disease_setting,trial_keyword1, trial_keyword_2, pt_genes, pt_stage, trial_stages, date_screened in results:
    url = 'https://www.cancer.gov/about-cancer/treatment/clinical-trials/search/v?' +nci_number
    writer.writerow([pt_number, type, therapy_type, mrn, pat_name, nci_number, nct_number, early_stage_resectable, early_stage_unresectable, advanced_first_line,advanced_second_line, title,spanish_title,english_description, spanish_description, trial_keyword1, trial_keyword_2, pt_genes, pt_stage, trial_stages, date_screened,disease_setting])

outfile.close()

#all of below has been moved to pretty

#mrns = dict()

#with open("matches/matches_" + today.strftime("%Y-%m-%d") + ".txt", newline='') as csvfile:
#    reader = csv.DictReader(csvfile)
#    for row in reader:
#        if(row['mrn'] in mrns):
#            mrns[row['mrn']].append(row)
#        else:
#            mrns[row['mrn']] = [row]

#for mrn in mrns:
#    doc = docx.Document()
#    thistype = -1
#    p = doc.add_paragraph()
#    run = p.add_run("MRN: " + mrn)
#    p = doc.add_paragraph()
#    run = p.add_run("PT Stage: " + mrns[mrn][0]['pt_stage'])
#    p = doc.add_paragraph()
#    run = p.add_run("PT Genes: " + mrns[mrn][0]['pt_genes'])    
#    
#    for row in mrns[mrn]:
#        if(row['type'] != thistype):
#            p = doc.add_paragraph()
#            thistype = row['type']
#            if(row['type'] == '1'):
#                run = p.add_run("Tier " + row['type'] + " matches, targeted therapy")    
#            elif (row['type'] == '2'):
#                run = p.add_run("Tier " + row['type'] + " matches, immunotherapy")
#            elif (row['type'] == '3'):
#                run = p.add_run("Tier " + row['type'] + " matches, antibody/drug conjugate")
#            elif (row['type'] == '4'):
#                run = p.add_run("Tier " + row['type'] + " matches, all others")
#                
#            
#        p = doc.add_paragraph()
#        #run = p.add_run("Study Number: " + row['nct_number'])           
#        table = doc.add_table(rows=2, cols=3)
#        r = table.rows[0].cells
#        r[0].text = "Study Number: " + row['nct_number']
#        r[1].text = "Study Name: " + row['title']
#        r[2].text = "Therapy Type: " + row['type']
#        r = table.rows[1].cells
#        if(row['type'] == '1'):
#            r[2].text = "Basis of Match(genes): " + row['trial_keyword1'] + ' ' + row['trial_keyword_2']
#        
#    doc.save(mrn + ".docx")
