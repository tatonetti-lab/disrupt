
import docx
import docx.document


def add_intro_text(doc : docx.document.Document, doc_type ):

    if doc_type == "PT_SPANISH":
        p_header = doc.add_paragraph()
        title = p_header.add_run("DISRUPT : NOTIFICACIÓN AL PACIENTE DE LOS RESULTADOS DE LA CONEXIÓN A ENSAYOS CLINICOS")
        title.bold = True

        p = doc.add_paragraph()

        title = p.add_run("Cómo hicimos este informe")
        title.bold = True
        title.add_break()
        p.add_run("Analizamos la información dada por su médico sobre su cáncer. Esto incluía detalles sobre los cambios en el tumor, cómo su sistema inmunitario está lidiando con el tumor y el estadio del tumor. Utilizamos esta información para hacer una lista de estudios de investigación en Columbia y otros lugares que podrían ser útiles para usted y su tratamiento.")

        p2 = doc.add_paragraph()
        title = p2.add_run("Cómo utilizar este informe ")
        title.bold = True
        title.add_break()
        p2.add_run("Este informe se hizo para incluir todos los ensayos a los que podría unirse. Es posible que haya muchos que realmente no se apliquen a usted, incluso los enumeramos. Utilizamos nuestro sistema para organizar ensayos en función de cuáles tienen más probabilidades de ser adecuados para usted. Dividimos estos ensayos en diferentes grupos:")

        doc.add_paragraph("Compatibilidad excelente: hay un cambio en uno o más genes (llamado mutación) en el tumor y hay un ensayo para personas con la misma mutación.",
                        style='List Bullet')
        doc.add_paragraph("Compatibilidad Buena: se está probando un nuevo tipo de tratamiento para su tipo de tumor, pero es posible que no se pruebe para el cambio específico en los genes del tumor (mutación).",
                        style='List Bullet')
        doc.add_paragraph("Compatibilidad Posible: No estamos seguros de si este ensayo es adecuado para su tumor, pero queremos que usted y su médico lo sepan por si acaso.",
                        style='List Bullet')
        doc.add_paragraph("Otra compatibilidad: Una posible compatibilidad que nos cuesta clasificar.",
                        style='List Bullet')

        doc.add_paragraph("Incluso si un estudio parece ser una buena opción para usted, puede haber otras razones por las que no es la opción correcta. Por ejemplo, podría ser una buena opción para las personas con cáncer de pulmón y su mutación genética específica, pero solo para aquellas que ya se han sometido a una cirugía, y es posible que su médico piense que usted no es candidato/a para la cirugía. La mejor manera de usar este informe es hablar sobre cualquier cosa que le interese con su médico, quien también recibirá una copia de este informe.")
        
        p3 = doc.add_paragraph()
        title = p3.add_run("¿Qué sucede si tengo preguntas sobre este informe?")
        title.bold = True
        title.add_break()
        p3.add_run("La mejor persona para preguntar sobre esto es su médico. Hacemos todo lo posible para proporcionar información precisa, pero puede haber errores. Su médico es la mejor persona para confirmar esto. Además, la información puede cambiar con el tiempo, incluidas las versiones de pruebas disponibles. Hacemos todo lo posible para mantener esta información actualizada, pero es posible que su médico conozca ensayos que no se mencionan aquí.")

        p4 = doc.add_paragraph()
        title = p4.add_run("Cómo unirse a un ensayo clínico")
        title.bold = True
        title.add_break()
        p4.add_run("Incluso si encontramos un buen ensayo para usted, es posible que tenga que cumplir con ciertos requisitos para participar en el ensayo y recibir el tratamiento. Cada ensayo es diferente, por lo que es importante hablar con su médico al respecto. Si bien hemos tratado de hacer que la información sea fácil de entender, siempre es mejor preguntarle a su médico si tiene alguna pregunta. Su médico puede proporcionarle la información más precisa y mantenerlo/la actualizado/a sobre cualquier nuevo desarrollo o ensayo que esté disponible con el tiempo.").add_break()
      
        return doc
    
    else: 
        p_header = doc.add_paragraph()
        title = p_header.add_run("DISRUPT :  PATIENT NOTIFICATION OF TRIAL MATCHING RESULTS")
        title.bold = True

        p = doc.add_paragraph()

        title = p.add_run("How we made this report")
        title.bold = True
        title.add_break()
        p.add_run("We looked at the information from your doctor about your cancer. This included details about the changes in your tumor, how your immune system is dealing with the tumor, and the stage of your tumor. We used this information to make a list of research studies at Columbia and other places that might be helpful for you and your treatment.")

        p2 = doc.add_paragraph()
        title = p2.add_run("How to use this report")
        title.bold = True
        title.add_break()
        p2.add_run("This report was made to include all the trials you might be able to join. There might be many that don't really apply to you even if we list them. We use our system to organize trials based on which ones are most likely to be right for you. We put these matches into different groups:")

        doc.add_paragraph("Excellent Match: You have a change in one or more genes (called a mutation) in your tumor and there's a trial for people with the same mutation.",
                        style='List Bullet')
        doc.add_paragraph("Good Match: A new kind of treatment is being tested for your type of tumor, but it might not be tested for the specific change in your tumor genes (mutation).",
                        style='List Bullet')
        doc.add_paragraph("Possible Match: We're not sure if this trial is right for your tumor, but we want you and your doctor to know about it just in case.",
                        style='List Bullet')
        doc.add_paragraph("Other Match: A possible match that is hard for us to classify",
                        style='List Bullet')

        doc.add_paragraph("Even if a study seems like it could be a good fit for you, there might be other reasons why it's not the right choice. For example, it might be a good match for people with lung cancer and your specific genetic mutation, but only for those who have already had surgery, and your doctor might not think you're a candidate for surgery. The best way to use this report is to talk about anything you're curious about with your doctor, who will also receive a copy of this report.")
        
        p3 = doc.add_paragraph()
        title = p3.add_run("What if I have questions about this report")
        title.bold = True
        title.add_break()
        p3.add_run("The best person to ask about this is your doctor. We do our best to provide accurate information, but there might be mistakes. Your doctor is the best person to confirm this. Also, information can change over time, including available trials. We do our best to keep this information up to date, but your doctor may know about trials not listed here.")

        p4 = doc.add_paragraph()
        title = p4.add_run("How to join a trial")
        title.bold = True
        title.add_break()
        p4.add_run("Even if we find a good trial for you, you may need to meet certain requirements to be part of the trial and receive the treatment. Each trial is different, so it's important to talk to your doctor about this. While we've tried to make the information easy to understand, it's always best to ask your doctor if you have any questions. Your doctor can provide you with the most accurate information and keep you updated on any new developments or trials that become available over time.").add_break()

        
        return doc