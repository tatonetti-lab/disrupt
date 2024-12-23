import docx
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

import json
import csv
from datetime import datetime

import docx.shared

from report_intro_text import add_intro_text

today = datetime.today().strftime("%Y-%m-%d")
#today = "2024-11-21"
mrns = dict()


# func to add bg color to table cells
def set_bg_color(cell, color):
    """
    set background shading for Header Rows
    """
    tblCell = cell._tc
    tblCellProperties = tblCell.get_or_add_tcPr()
    clShading = OxmlElement('w:shd')
    clShading.set(qn('w:fill'), color) 
    tblCellProperties.append(clShading)
    return cell

# open matchlist and format with mrn as key and value as pt and trial info
# pt[studies] is an array of all trials 
with open("matches/matches_" + today + ".txt", newline='') as csvfile:
    reader = csv.DictReader(csvfile)
    for row in reader:
        temp_row = row.copy()
        temp_row.pop('mrn')
        temp_row.pop('pt_stage')

        pt_genes = [i.strip() for i in temp_row.pop('pt_genes').split(",")]

        if(row['mrn'] in mrns):
            mrns[row['mrn']]['studies'].append( temp_row )
        else:
            mrns[row['mrn']] = {
                "mrn": row['mrn'],
                "pt_name": row["patient_name"],
                "pt_stage": row['pt_stage'],
                "pt_genes": set(),
                "disease_setting": row["disease setting"],
                "studies": [temp_row]
            }

        if pt_genes != ['']: 
            mrns[row['mrn']]['pt_genes'].update( pt_genes ) # only update if not empty



mrns = list(mrns.values()) # create values into a list
# mrn = [{ mrn, pt name, stage, disease, [genes], [trials]}, ...]


pt_stats = {}

doc_types = ['DR', "PT_ENGLISH", "PT_SPANISH"]

for pt in mrns:

    pt_stats[pt['mrn']] = {
        "name": pt['pt_name'],
        "Excellent": 0,
        "Good": 0,
        "Possible": 0,
        "Other": 0
    }

    stat_curr = pt_stats[pt['mrn']]

    for row in pt['studies']:
        if row['type'] == "1":
            stat_curr["Excellent"]+=1
        elif row['type'] == "2":
            stat_curr["Good"]+=1
        elif row['type'] == "3":
            stat_curr["Possible"]+=1
        else:
            stat_curr["Other"]+=1


    for doc_type in doc_types:
        doc = docx.Document()

        thistype = -1

        # title: mrn name stage disease setting genes
        p = doc.add_paragraph()

        pt_name = p.add_run(pt['pt_name'])
        pt_name.font.size = Pt(16)

        pt_name.add_break()

        mrn_title = p.add_run("MRN: ")
        mrn_title.bold = True

        mrn_num = p.add_run(pt['mrn'])
        mrn_num.font.size = Pt(14)

        mrn_num.add_break()

        stage_title = p.add_run("Stage: ")
        stage_title.bold = True

        if pt['pt_stage'] != '':
            stage = p.add_run(pt['pt_stage'])
            stage.font.size = Pt(14)
            stage.add_break()
        else:
            p.add_run("N/A").add_break()
        

        dis_set_title = p.add_run("Disease Setting: ")
        dis_set_title.bold = True
        if pt['disease_setting'] != '':
            dis_set = p.add_run(pt['disease_setting'])
            dis_set.font.size = Pt(14)
        else:
            p.add_run("N/A")

        # gene
        p = doc.add_paragraph()
        gene_title = p.add_run("Relevant Biomarkers:")
        gene_title.bold = True
        gene_title.add_break()

        for gene in sorted(pt['pt_genes']):
            genes = p.add_run(gene)
            genes.font.size = docx.shared.Pt(16)
            genes.add_break()
        
        
        if len(pt['pt_genes']) == 0:
            p.add_run("N/A").add_break()

        
        doc = add_intro_text(doc, doc_type)

        p_studies_title = doc.add_paragraph()
        p_studies_title.add_run("POTENTIAL TRIAL MATCHES").bold = True
        
        # iterate through studies
        for row in pt['studies']:

            #if (doc_type == "PT_SPANISH"):
            #    if (row['spanish title'] == ""):
            #        continue

            """
            if(row['type'] != thistype):
                p_type = doc.add_paragraph()
                thistype = row['type']
                #p_type.style.font.italic = True
                if(row['type'] == '1'):
                    run = p_type.add_run("Tier " + row['type'] + ": Targeted Therapy")    
                elif (row['type'] == '2'):
                    run = p_type.add_run("Tier " + row['type'] + ": Immunotherapy")
                elif (row['type'] == '3'):
                    run = p_type.add_run("Tier " + row['type'] + ": Antibody/Drug Conjugate")
                elif (row['type'] == '4'):
                    run = p_type.add_run("Tier " + row['type'] + ": All Others")
            """
            
            # TABLE

            # row for description not included in dr version
            table_rows = 4
            if doc_type == "DR":
                table_rows=3

            table = doc.add_table(rows=table_rows, cols=3)

            # row 1: (nct and therapy type / resectable stuff)
            r = table.rows[0].cells
            
            r[0].text = ""
            nct = r[0].paragraphs[0].add_run(row['nct_number'])
            nct.bold = True
            nct.font.size = Pt(16)
            

            r[1].text = ""
            tt = r[1].paragraphs[0].add_run("THERAPY TYPE: ") 
            tt.bold = True
            tt.font.size = Pt(10)
            r[1].paragraphs[0].add_run(row['therapy_type'])

            # early_stage_resectable, early_stage_unresectable, advanced_first_line, advanced_second_line
            therapy_stages = []

            if row['early_stage_resectable'] == "Y":
                therapy_stages.append("Early Stage Resectable")
            if row['early_stage_unresectable'] == "Y":
                therapy_stages.append("Early Stage Unresectable")
            if row['advanced_first_line'] == "Y":
                therapy_stages.append("Advanced First Line")
            if row['advanced_second_line'] == "Y":
                therapy_stages.append("Advanced Second Line")
            
            """
            if len(therapy_stages) != 0:
                r[1].paragraphs[0].add_run().add_break()
                r[1].paragraphs[0].add_run( ", ".join(therapy_stages) )
            """

            r[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            r[1].merge(r[2])

            set_bg_color(r[0], 'cce0eb')
            set_bg_color(r[1], 'cce0eb')
            set_bg_color(r[2], 'cce0eb')

            # row 2: title of study
            
            r = table.rows[1].cells
            r[0].text = "" 
            if doc_type == "PT_SPANISH" and row['spanish title'] != "":    
                r[0].paragraphs[0].add_run(row['spanish title']).italic = True
            else:
                r[0].paragraphs[0].add_run(row['title']).italic = True

            r[0].merge(r[1])
            r[0].merge(r[2])

            # row 3: match reason and trial disease stages

            r = table.rows[2].cells

            r[0].text = ""
            reason = r[0].paragraphs[0].add_run('REASON FOR MATCH: ')
            reason.bold = True
            reason.font.size = Pt(10)
            if (row['trial_keyword1'] == "N/A" or row['trial_keyword1'] == ""):
                r[0].paragraphs[0].add_run("N/A")
            else:
                r[0].paragraphs[0].add_run(row['trial_keyword1'] + ' ' + row['trial_keyword_2']).font.size = Pt(14)

            r[0].paragraphs[0].add_run().add_break()

            strength = r[0].paragraphs[0].add_run("STRENGTH OF MATCH: ")
            strength.bold = True
            strength.font.size = Pt(10)

            if row['type'] == "1":
                r[0].paragraphs[0].add_run("Excellent")
              
            elif row['type'] == "2":
                r[0].paragraphs[0].add_run("Good")
                
            elif row['type'] == "3":
                r[0].paragraphs[0].add_run("Possible")
                
            else:
                r[0].paragraphs[0].add_run("Other")
                

            r[0].paragraphs[0].add_run().add_break()

            dis_stages = r[0].paragraphs[0].add_run("DISEASE STAGES CONSIDERED: ")
            dis_stages.bold = True
            dis_stages.font.size = Pt(10)
            if len(therapy_stages) == 0:
                r[0].paragraphs[0].add_run("N/A")
            else:
                r[0].paragraphs[0].add_run(", ".join(therapy_stages)).font.size = Pt(12)

            r[0].merge(r[1])
            r[0].merge(r[2])

            # row 4: description IF not dr format

            if doc_type != "DR":

                r = table.rows[3].cells
                r[0].text = "" 
                if doc_type == "PT_SPANISH":       
                    r[0].paragraphs[0].add_run(row['spanish description'])
                else:
                    r[0].paragraphs[0].add_run(row['english description'])
                r[0].merge(r[1])
                r[0].merge(r[2])

                set_bg_color(r[0], 'f7f7f7')
                set_bg_color(r[1], 'f7f7f7')
                set_bg_color(r[2], 'f7f7f7')

            doc.add_paragraph()

        
        # add borders, center align, line spacing, overall font

        for t in doc.tables:
            tbl = t._tbl

            for cell in tbl.iter_tcs():
                tcPr = cell.tcPr
                tcBorders = OxmlElement("w:tcBorders")

                left = OxmlElement("w:left")
                left.set(qn("w:val"), "none")

                right = OxmlElement("w:right")
                right.set(qn("w:val"), "none")

                top = OxmlElement("w:top")
                top.set(qn("w:val"), "none")

                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "4")
                bottom.set(qn("w:color"), "#c2c2c2")

                tcBorders.append(top)
                tcBorders.append(left)
                tcBorders.append(bottom)
                tcBorders.append(right)
                tcPr.append(tcBorders)

            for r in range(len(table.rows)):
                for c in range(len(table.columns)):
                    t.cell(r,c).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    pt_num = 3
                    t.cell(r,c).paragraphs[0].paragraph_format.space_after = Pt(pt_num)
                    t.cell(r,c).paragraphs[0].paragraph_format.space_before = Pt(pt_num)

        for paragraph in doc.paragraphs:
            paragraph.style.font.name = "Aptos"
            paragraph.style.font.size = Pt(12)
        
    
            
        doc.save("matches/" + pt['mrn'] + "_matchlist_" + doc_type + ".docx")


new_file = open("matches/pt_match_stats.csv", 'w', newline='')
csv_write = csv.writer(new_file, delimiter=",")

csv_write.writerow(["Count of Matches per Patient"])
csv_write.writerow([])

i = 0
for pt in pt_stats:
    csv_write.writerow([pt + ": " + pt_stats[pt]["name"]])
    csv_write.writerow(["TOTAL", pt_stats[pt]["Excellent"] + pt_stats[pt]["Good"] + pt_stats[pt]["Possible"] + pt_stats[pt]["Other"]])
    csv_write.writerow(["Excellent", pt_stats[pt]["Excellent"]])
    csv_write.writerow(["Good", pt_stats[pt]["Good"]])
    csv_write.writerow(["Possible", pt_stats[pt]["Possible"]])
    csv_write.writerow(["Other", pt_stats[pt]["Other"]])
    csv_write.writerow([])

new_file.close()
   